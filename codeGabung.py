import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def read_and_filter_csv(file_path):
    """Reads a CSV file and returns only numeric values in a DataFrame."""
    df = pd.read_csv(file_path)
    return df

def combine_csv_files(file_paths_kf, file_paths_non_kf):
    """Combines multiple sets of CSV files into one DataFrame with specific columns."""
    combined_kf_df = pd.DataFrame()
    combined_non_kf_df = pd.DataFrame()
    
    for file_path in file_paths_kf:
        df = read_and_filter_csv(file_path)
        combined_kf_df = pd.concat([combined_kf_df, df], ignore_index=True)
    
    for file_path in file_paths_non_kf:
        df = read_and_filter_csv(file_path)
        combined_non_kf_df = pd.concat([combined_non_kf_df, df], ignore_index=True)
    
    combined_df = pd.DataFrame()
    combined_df['Tanpa Kalman'] = combined_non_kf_df.squeeze().round(3)
    combined_df['Dengan Kalman'] = combined_kf_df.squeeze().round(3)
    
    return combined_df

def set_table_border(table):
    """Set borders for the table."""
    tbl = table._element
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

def save_to_word(df, output_file):
    """Saves the DataFrame to a Word document with table borders."""
    document = Document()
    document.add_heading('Combined Data', 0)

    table = document.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
    
    set_table_border(table)

    document.save(output_file)
    print(f"Word document saved to {output_file}")

# Example usage
file_paths_kf = [
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\1\outputErrorTitikGaris(KF).csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\2\outputErrorTitikGaris(KF).csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\3\outputErrorTitikGaris(KF).csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\4\outputErrorTitikGaris(KF).csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\5\outputErrorTitikGaris(KF).csv'
]  # Replace with your KF CSV file paths

file_paths_non_kf = [
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\1\outputErrorTitikGaris.csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\2\outputErrorTitikGaris.csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\3\outputErrorTitikGaris.csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\4\outputErrorTitikGaris.csv',
    r'Data Penelitian\Data Final v4\Dinamis\Ruang 2\Ruang Isi\Posisi 1\5\outputErrorTitikGaris.csv',
]  # Replace with your non-KF CSV file paths

output_word_file = 'combined_output.docx'

combined_df = combine_csv_files(file_paths_kf, file_paths_non_kf)
save_to_word(combined_df, output_word_file)
